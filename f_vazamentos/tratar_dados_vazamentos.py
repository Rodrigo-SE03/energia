import zipfile
import os
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from PIL import Image, ImageOps
import re

#Função para descompactar o arquivo zip
def tratar_dados(dados_empresa,form_vazamentos):
    dados_empresa['Empresa'] = form_vazamentos.empresa.data
    dados_empresa['CNPJ'] = form_vazamentos.cnpj.data
    dados_empresa['Endereço'] = form_vazamentos.endereco.data
    dados_empresa['Contato'] = form_vazamentos.contato_nome.data
    dados_empresa['Departamento'] = form_vazamentos.contato_depto.data
    dados_empresa['E-mail'] = form_vazamentos.contato_email.data
    dados_empresa['Telefone'] = form_vazamentos.contato_fone.data
    dados_empresa['RT'] = form_vazamentos.rt.data
    dados_empresa['Membro 1'] = form_vazamentos.e1.data
    dados_empresa['Membro 2'] = form_vazamentos.e2.data
    dados_empresa['Membro 3'] = form_vazamentos.e3.data
    dados_empresa['Membro 4'] = form_vazamentos.e4.data
    dados_empresa['Setores'] = form_vazamentos.setores.data

    return dados_empresa
#--------------------------------------------------------------------------------------------------------

#Função para descompactar o arquivo zip
def unzip(file,folder):
    with zipfile.ZipFile(file, 'r') as zip_ref:
        zip_ref.extractall(f'{folder}')
    os.remove(f'{folder}/{file.filename}')
#--------------------------------------------------------------------------------------------------------

#Função para inserir os pontos de fuga registrados no modelo do relatório preliminar
def relatorio(file,folder,dados_empresa):
    doc = Document(f'{folder}/vazamentos_modelo.docx')
    
    last = []
    id = 0
    cont = 0
    idx=10000
    for p in doc.paragraphs:
        if '[}(--O_O--){]' in p.text:
            idx = id
            p.text = p.text.replace('[}(--O_O--){]','\n')
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if id > idx:
            last.append(p)
        id+=1
    
    folder = f'{folder}/{file.split('.')[0]}'
    files = os.listdir(folder)
    img_format = files[0].split('.')[1]
    os.mkdir(f'{folder}/Fotos')
    for file in files:
        if file != 'Vazamentos.xlsx':
            new_img = Image.open(f'{folder}/{file}')
            new_img = ImageOps.exif_transpose(new_img)
            new_img.save(f'{folder}/Fotos/{file}')
            os.remove(f'{folder}/{file}')

    df = pd.read_excel(open(f'{folder}/Vazamentos.xlsx','rb'))
    document = doc
    qtd_total = 0
    qtd_pequeno = 0
    qtd_medio = 0
    qtd_grande = 0
    qtd_extragrande = 0

    for index,row in df.iterrows():
        table = document.add_table(rows = 0,cols = 3,style = "Table Grid")
        pic = f"{folder}/Fotos/{row['Foto']}.{img_format}"
        itens = {
                'Local':row["Local"],
                'Elemento/Componente':row["Componente"],
                'Classificação':row["Classificação"],
                'Data':row['Data'],
                'Quantidade':row["Quantidade"],
                'Observações':row["Observações"]
                }
        if type(itens['Observações']) != str:
            itens['Observações'] = ""
        
        data = list(re.search(r'([0-9]+)[\/]([0-9]+)[\/]([0-9]+)',str(itens['Data'])).groups())

        qtd_total = qtd_total + itens['Quantidade']
        if itens['Classificação'] == "Pequeno":
            qtd_pequeno = qtd_pequeno + itens['Quantidade']
        elif itens['Classificação'] == "Médio":
            qtd_medio = qtd_medio + itens['Quantidade']
        elif itens['Classificação'] == "Grande":
            qtd_grande = qtd_grande + itens['Quantidade']
        else:
            qtd_extragrande = qtd_extragrande + itens['Quantidade']
        if int(data[1]) < 10:
            data[1] = f'0{data[1]}'
        itens['Data'] = f"{data[0]}/{data[1]}/{data[2]}"
        
        for item in itens.items():
            
            row_cells = table.add_row().cells
            p1 = row_cells[1].paragraphs[0]
            paragraph_format = p1.paragraph_format
            paragraph_format.line_spacing = 2.0
            p1.add_run(str(item[0])).bold = True

            paragraph_format.space_after
            p2 = row_cells[2].paragraphs[0]
            paragraph_format = p2.paragraph_format
            paragraph_format.line_spacing = 1.0
            p2.add_run(str(item[1]))
            #row_cells[2].text = str(item[1])

        i = 1
        photo_col = table.column_cells(0)
        photo_cell = photo_col[0]
        while i<6:
            photo_cell.merge(photo_col[i])
            i+=1
        photo_cell.text = ""
        paragraph = photo_cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(pic,1300000)

        i=0
        while i<6:
            j=0
            while j<3:
                table.cell(i,j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if j == 0:
                    table.cell(i,j).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                j+=1
            i+=1
        
        p = document.add_paragraph()
        


        if cont == 2:
            p.add_run("\n\n\n\n\n")
            cont = -1
        cont+=1

    p_last = p
    for p in last:
        new_p = p
        para_last = p_last._p
        para_last.addnext(new_p._p)
        p_last = new_p

    def analise_custos(qtd_total,tarifa_p_ml,tarifa_p_c,tarifa_fp_ml,tarifa_fp_c,hp,hfp,dias_p_mes):
        tp = tarifa_p_c+tarifa_p_ml
        tfp = tarifa_fp_ml+tarifa_fp_c
        desp_energia = qtd_total*0.3*(hp+hfp)*dias_p_mes*12
        custo = qtd_total*0.3*(hp*tp+hfp*tfp)*dias_p_mes*12

        return[desp_energia,custo]

    #print(analise_custos(qtd_total,1.56898,0.85865,0.148332,0.14833,2,14,22))

    for p in doc.paragraphs:
        if '[-EMPRESA-]' in p.text:
            p.text = p.text.replace('[-EMPRESA-]',dados_empresa['Empresa'])
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
        elif '[-CNPJ-]' in p.text:
            p.text = p.text.replace('[-CNPJ-]',f'CNPJ: {dados_empresa['CNPJ']}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
        elif '[-ENDEREÇO 1-]' in p.text:
            p.text = p.text.replace('[-ENDEREÇO 1-]',f'{dados_empresa['Endereço'].split(';')[0]}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
        elif '[-ENDEREÇO 2-]' in p.text:
            p.text = p.text.replace('[-ENDEREÇO 2-]',f'{dados_empresa['Endereço'].split(';')[1]}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)  
        elif '[-ENDEREÇO 3-]' in p.text:
            p.text = p.text.replace('[-ENDEREÇO 3-]',f'{dados_empresa['Endereço'].split(';')[2]}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)    
        elif '[-CONTATO-]' in p.text:
            p.text = p.text.replace('[-CONTATO-]',f'Contato: {dados_empresa['Contato']}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12) 
                run.font.bold = True
        elif '[-DEPARTAMENTO-]' in p.text:
            p.text = p.text.replace('[-DEPARTAMENTO-]',f'Departamento: {dados_empresa['Departamento']}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12) 
        elif '[-EMAIL-]' in p.text:
            p.text = p.text.replace('[-EMAIL-]',f'E-mail: {dados_empresa['E-mail']}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
        elif '[-TELEFONE-]' in p.text:
            p.text = p.text.replace('[-TELEFONE-]',f'Telefone: {dados_empresa['Telefone']}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12) 
        elif '[-RT-]' in p.text:
            p.text = p.text.replace('[-RT-]',f'{dados_empresa['RT']}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12) 
        elif '[-MEMBRO 1-]' in p.text:
            p.text = p.text.replace('[-MEMBRO 1-]',f'{dados_empresa['Membro 1']}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12) 
        elif '[-MEMBRO 2-]' in p.text:
            p.text = p.text.replace('[-MEMBRO 2-]',f'{dados_empresa['Membro 2']}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12) 
        elif '[-MEMBRO 3-]' in p.text:
            p.text = p.text.replace('[-MEMBRO 3-]',f'{dados_empresa['Membro 3']}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12) 
        elif '[-MEMBRO 4-]' in p.text:
            p.text = p.text.replace('[-MEMBRO 4-]',f'{dados_empresa['Membro 4']}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12) 
        elif '[-VP-]' in p.text:
            p.text = p.text.replace('[-VP-]',str(qtd_pequeno))
            p.text = p.text.replace('[-VM-]',str(qtd_medio))
            p.text = p.text.replace('[-VG-]',str(qtd_grande))
            p.text = p.text.replace('[-VE-]',str(qtd_extragrande))
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
        elif '[-TOTAL-]' in p.text:
            p.text = p.text.replace('[-TOTAL-]',str(qtd_total))
            p.text = p.text.replace('[-SETORES-]',dados_empresa['Setores'])   
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)

    print(f"Total: {qtd_total}\nPequenos: {qtd_pequeno}\nMedios:{qtd_medio}\nGrandes:{qtd_grande}\nExtragrandes: {qtd_extragrande}")
    document.save(f'{folder}/tabelas.docx')
#--------------------------------------------------------------------------------------------------------